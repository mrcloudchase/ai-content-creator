import docx
from typing import Dict, List, Any, Optional
import io
import os
import re
import sys
import json
from app.services.token_service import TokenService
from app.config.settings import OpenAISettings

class DocxParserError(Exception):
    """Custom exception for document parsing errors"""
    pass

class TokenLimitError(DocxParserError):
    """Exception raised when document exceeds token limit"""
    pass

class DocxParser:
    """Service for parsing .docx documents"""
    
    @staticmethod
    def parse_document(file_content: bytes) -> Dict[str, Any]:
        """
        Parse a .docx document and maintain its structure
        
        Args:
            file_content: Binary content of the .docx file
            
        Returns:
            Dictionary containing the parsed document structure
            
        Raises:
            DocxParserError: If there's an error parsing the document
        """
        # Input validation
        assert file_content is not None, "File content cannot be None"
        assert isinstance(file_content, bytes), "File content must be bytes"
        assert len(file_content) > 0, "File content cannot be empty"
        
        try:
            # Load the document from bytes
            doc = docx.Document(io.BytesIO(file_content))
            
            # Parse document content
            result = {
                "title": DocxParser._get_title(doc),
                "paragraphs": DocxParser._get_paragraphs(doc),
                "tables": DocxParser._get_tables(doc),
                "headings": DocxParser._get_headings(doc),
                "metadata": DocxParser._get_metadata(doc),
            }
            
            # Output validation
            assert isinstance(result, dict), "Result must be a dictionary"
            assert "title" in result, "Result must have a title"
            assert "paragraphs" in result, "Result must have paragraphs"
            assert "tables" in result, "Result must have tables"
            assert "headings" in result, "Result must have headings"
            assert "metadata" in result, "Result must have metadata"
            
            return result
        except AssertionError as e:
            # Re-raise assertion errors for debugging
            raise
        except Exception as e:
            raise DocxParserError(f"Error parsing document: {str(e)}")
    
    @staticmethod
    def extract_text(file_content: bytes) -> str:
        """
        Extract text content from a .docx document, ensuring the output is JSON-compatible
        
        Args:
            file_content: Binary content of the .docx file
            
        Returns:
            String containing all text from the document, sanitized for JSON compatibility
            
        Raises:
            DocxParserError: If there's an error parsing the document
            TokenLimitError: If the document exceeds the configured token limit
        """
        # Input validation
        assert file_content is not None, "File content cannot be None"
        assert isinstance(file_content, bytes), "File content must be bytes"
        assert len(file_content) > 0, "File content cannot be empty"
        
        try:
            # Load the document from bytes
            doc = docx.Document(io.BytesIO(file_content))
            
            # Extract text from paragraphs, preserving the document flow
            text_parts = []
            
            # Process document elements in sequential order
            for element in doc.element.body.iter():
                # Skip the overall body element
                if element.tag.endswith('body'):
                    continue
                
                # Process paragraphs (includes headings, normal paragraphs, and list items)
                if element.tag.endswith('p'):
                    paragraph = None
                    
                    # Find the corresponding paragraph object
                    for p in doc.paragraphs:
                        if p._element is element:
                            paragraph = p
                            break
                    
                    if paragraph and paragraph.text.strip():
                        # Check if it's a list item (has a numPr or bullet/number character)
                        is_list_item = False
                        list_format = ""
                        
                        # Look for numbering properties
                        num_pr = element.xpath('.//w:numPr')
                        if num_pr:
                            is_list_item = True
                            # Approximation of indentation level based on list item
                            ilvl = element.xpath('.//w:ilvl/@w:val')
                            indent = int(ilvl[0]) if ilvl else 0
                            list_format = "  " * indent + "• "
                        
                        # Add the paragraph with appropriate formatting
                        if is_list_item:
                            text_parts.append(f"{list_format}{paragraph.text}")
                        else:
                            text_parts.append(paragraph.text)
                
                # Process tables
                elif element.tag.endswith('tbl'):
                    table = None
                    
                    # Find the corresponding table object
                    for t in doc.tables:
                        if t._element is element:
                            table = t
                            break
                    
                    if table:
                        table_text_parts = []
                        for row in table.rows:
                            row_texts = []
                            for cell in row.cells:
                                row_texts.append(cell.text.strip())
                            table_text_parts.append(" | ".join(row_texts))
                        
                        # Add the table rows as separate lines
                        text_parts.extend(table_text_parts)
            
            # Join all text parts with appropriate spacing
            # Simulate the way Word adds spacing between different elements
            text = "\n\n".join(text_parts)
            
            # Clean up excessive whitespace
            text = re.sub(r'\n{3,}', '\n\n', text)
            
            # Apply JSON compatibility processing
            text = DocxParser._make_json_compatible(text)
            
            # Verify JSON serialization works for both direct and nested scenarios
            try:
                # Test 1: Basic JSON serialization
                test_json = json.dumps({"prompt": text})
                
                # Test 2: Nested JSON scenario (simulating user copy-paste)
                nested_json = json.dumps({"outer": json.dumps({"prompt": text})})
                json.loads(nested_json)  # Make sure it parses correctly
                
                # Test 3: Direct usage in a prompt field (most common use case)
                ai_json = '{{"prompt": {}}}'.format(json.dumps(text))
                json.loads(ai_json)  # Should parse without errors
                
            except Exception as e:
                raise DocxParserError(f"Generated text is not JSON-compatible: {str(e)}")
            
            # Check token count against OpenAI max tokens
            try:
                # Get the max tokens from settings
                settings = OpenAISettings()
                token_result = TokenService.count_tokens(text, settings.default_model)
                
                # Define a more generous token limit for document inputs
                # This is separate from the OpenAI completion output limit
                DOCUMENT_TOKEN_LIMIT = 8192  # Allow documents up to 8k tokens
                
                # Verify token count is within limits
                if token_result["token_count"] > DOCUMENT_TOKEN_LIMIT:
                    raise TokenLimitError(
                        f"Document exceeds token limit: {token_result['token_count']} tokens "
                        f"(limit: {DOCUMENT_TOKEN_LIMIT}). Use the token counting endpoint "
                        f"for more information."
                    )
            except Exception as e:
                if isinstance(e, TokenLimitError):
                    raise
                # If there's another error with token counting, log it but don't block processing
                print(f"Warning: Could not verify token count: {str(e)}")
            
            # Output validation - empty documents just return empty string
            assert isinstance(text, str), "Result must be a string"
            
            # Return the text - even if empty
            return text
        except AssertionError as e:
            # Re-raise assertion errors for debugging
            raise
        except DocxParserError:
            # Re-raise our custom exceptions
            raise
        except Exception as e:
            raise DocxParserError(f"Error extracting text from document: {str(e)}")
    
    @staticmethod
    def _make_json_compatible(text: str) -> str:
        """
        Process text to ensure it's compatible with JSON serialization, even in nested JSON contexts
        
        Args:
            text: The original text to process
            
        Returns:
            Processed text that can be safely serialized in JSON and used in nested JSON objects
        """
        if not text:
            return text
            
        # Step 1: Normalize line breaks to \n
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        # Step 2: Remove problematic control characters (keep tabs and newlines)
        # This removes all control chars except \t (9), \n (10), and \r (13)
        text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
        
        # Step 3: Normalize whitespace (but preserve paragraph breaks)
        text = re.sub(r' {2,}', ' ', text)
        
        # Step 4: Replace Unicode quotes and apostrophes with ASCII ones
        # This is critical for JSON compatibility in nested contexts
        text = text.replace('"', '"').replace('"', '"')
        text = text.replace(''', "'").replace(''', "'")
        
        # Step 5: Replace other potentially problematic characters
        text = text.replace('—', '--').replace('–', '-')
        text = text.replace('…', '...')
        
        # Step 6: Ensure backslashes are properly escaped
        # Each backslash needs to be double-escaped in nested JSON contexts
        text = text.replace('\\', '\\\\')
        
        # Step 7: Special handling for quote characters in nested JSON contexts
        # For direct copy-paste into another JSON object
        text = text.replace('"', '\\"')
        
        return text
    
    @staticmethod
    def _get_title(doc: docx.Document) -> str:
        """Extract the document title"""
        if doc.paragraphs and doc.paragraphs[0].text:
            return doc.paragraphs[0].text
        return "Untitled Document"
    
    @staticmethod
    def _get_paragraphs(doc: docx.Document) -> List[Dict[str, Any]]:
        """Extract all paragraphs from the document"""
        paragraphs = []
        
        for para in doc.paragraphs:
            if not para.text.strip():
                continue
                
            # Skip if this is a heading (we'll handle these separately)
            if re.match(r'Heading \d+', para.style.name):
                continue
                
            paragraph_data = {
                "text": para.text,
                "style": para.style.name,
                "runs": [
                    {
                        "text": run.text,
                        "bold": run.bold,
                        "italic": run.italic,
                        "underline": run.underline,
                    }
                    for run in para.runs
                ]
            }
            paragraphs.append(paragraph_data)
            
        return paragraphs
    
    @staticmethod
    def _get_tables(doc: docx.Document) -> List[Dict[str, Any]]:
        """Extract all tables from the document"""
        tables = []
        
        for table in doc.tables:
            table_data = []
            
            for row in table.rows:
                row_data = []
                
                for cell in row.cells:
                    row_data.append(cell.text)
                    
                table_data.append(row_data)
                
            tables.append({"data": table_data})
            
        return tables
    
    @staticmethod
    def _get_headings(doc: docx.Document) -> List[Dict[str, Any]]:
        """Extract all headings from the document"""
        headings = []
        
        for para in doc.paragraphs:
            if re.match(r'Heading \d+', para.style.name):
                level = int(para.style.name.split()[-1])
                headings.append({
                    "text": para.text,
                    "level": level
                })
                
        return headings
    
    @staticmethod
    def _get_metadata(doc: docx.Document) -> Dict[str, str]:
        """Extract document metadata"""
        metadata = {}
        
        try:
            core_properties = doc.core_properties
            
            if core_properties.author:
                metadata["author"] = core_properties.author
            if core_properties.title:
                metadata["title"] = core_properties.title
            if core_properties.created:
                metadata["created"] = str(core_properties.created)
            if core_properties.modified:
                metadata["modified"] = str(core_properties.modified)
                
            return metadata
        except:
            # If metadata extraction fails, just return an empty dict
            return {} 