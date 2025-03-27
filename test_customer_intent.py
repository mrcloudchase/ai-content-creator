#!/usr/bin/env python3
import requests
import docx
import io
import json

def create_test_document():
    """Create a simple test document"""
    doc = docx.Document()
    
    # Add title
    doc.add_heading('User Authentication System PRD', 0)
    
    # Add overview
    doc.add_heading('1. Overview', level=1)
    doc.add_paragraph('This PRD outlines the requirements for implementing a secure user authentication system for our web application. The system will support email/password login, social auth, and 2FA.')
    
    # Add user stories
    doc.add_heading('2. User Stories', level=1)
    doc.add_paragraph('• As a user, I want to create an account so I can access personalized features', style='List Bullet')
    doc.add_paragraph('• As a user, I want to reset my password if I forget it', style='List Bullet')
    doc.add_paragraph('• As an admin, I want to manage user accounts to maintain security', style='List Bullet')
    
    # Save to bytes
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    return docx_bytes.getvalue()

def test_workflow():
    """Test the document parser -> customer intent workflow"""
    print("Creating test document...")
    doc_bytes = create_test_document()
    
    # Step 1: Parse the document
    print("\nStep 1: Sending document for parsing...")
    files = {"file": ("test.docx", doc_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    parse_response = requests.post("http://localhost:8000/api/v1/documents/extract-text", files=files)
    
    if parse_response.status_code != 200:
        print(f"Error parsing document: {parse_response.status_code}")
        print(parse_response.text)
        return
    
    print(f"✅ Document parsed successfully: {parse_response.status_code}")
    parsed_data = parse_response.json()
    
    # Step 2: Generate customer intent
    print("\nStep 2: Generating customer intent...")
    intent_request = {
        "document_text": parsed_data["document"],
        "max_tokens": 150,
        "temperature": 0.7
    }
    
    try:
        intent_response = requests.post("http://localhost:8000/api/v1/customer-intent/generate", json=intent_request)
        
        if intent_response.status_code != 200:
            print(f"❌ Error generating customer intent: {intent_response.status_code}")
            print(intent_response.text)
            return
        
        print(f"✅ Customer intent generated successfully: {intent_response.status_code}")
        intent_data = intent_response.json()
        
        print("\nCustomer Intent:")
        print(f"{'-'*50}\n{intent_data['intent']}\n{'-'*50}")
        print(f"\nModel: {intent_data['model']}")
        print(f"Token usage: {intent_data['usage']}")
    except Exception as e:
        print(f"❌ Error in customer intent request: {str(e)}")
        return
    
    # Step 3: Generate customer intent for a specific user type
    print("\nStep 3: Generating customer intent for an admin...")
    admin_intent_request = {
        "document_text": parsed_data["document"],
        "user_type": "admin",
        "max_tokens": 150,
        "temperature": 0.7
    }
    
    admin_intent_response = requests.post("http://localhost:8000/api/v1/customer-intent/generate", json=admin_intent_request)
    
    if admin_intent_response.status_code != 200:
        print(f"❌ Error generating admin customer intent: {admin_intent_response.status_code}")
        print(admin_intent_response.text)
        return
    
    print(f"✅ Admin customer intent generated successfully: {admin_intent_response.status_code}")
    admin_intent_data = admin_intent_response.json()
    
    print("\nAdmin Customer Intent:")
    print(f"{'-'*50}\n{admin_intent_data['intent']}\n{'-'*50}")

if __name__ == "__main__":
    test_workflow() 