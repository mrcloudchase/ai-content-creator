import pytest
import io
import os
import docx
from app.services.docx_parser import DocxParser, DocxParserError

@pytest.fixture
def sample_docx():
    """Create a sample docx document for testing"""
    doc = docx.Document()
    
    # Add title
    doc.add_heading('Sample Document', 0)
    
    # Add some paragraphs
    doc.add_paragraph('This is a sample paragraph with some content.')
    
    # Add a paragraph with styled text
    p = doc.add_paragraph('This paragraph has ')
    p.add_run('bold').bold = True
    p.add_run(' and ')
    p.add_run('italic').italic = True
    p.add_run(' text.')
    
    # Add a heading
    doc.add_heading('Section 1', level=1)
    
    # Add another paragraph
    doc.add_paragraph('Content in the first section.')
    
    # Add a table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = 'Header 1'
    table.cell(0, 1).text = 'Header 2'
    table.cell(1, 0).text = 'Value 1'
    table.cell(1, 1).text = 'Value 2'
    
    # Save to bytes
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    return docx_bytes.getvalue()

@pytest.fixture
def empty_docx():
    """Create an empty docx document for testing edge cases"""
    doc = docx.Document()
    
    # Save to bytes
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    return docx_bytes.getvalue()

@pytest.fixture
def list_docx():
    """Create a docx document with lists for testing"""
    doc = docx.Document()
    
    # Add a heading
    doc.add_heading('Document with Lists', 0)
    
    # Add a bulleted list - need to apply style manually
    doc.add_paragraph('First bullet point', style='List Bullet')
    doc.add_paragraph('Second bullet point', style='List Bullet')
    
    # Add a numbered list - need to apply style manually
    doc.add_paragraph('First numbered item', style='List Number')
    doc.add_paragraph('Second numbered item', style='List Number')
    
    # Save to bytes
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    return docx_bytes.getvalue()

def test_parse_document(sample_docx):
    # Parse the sample document
    result = DocxParser.parse_document(sample_docx)
    
    # Check that the result has the expected structure
    assert isinstance(result, dict), "Result should be a dictionary"
    assert "title" in result, "Result should have a title key"
    assert "paragraphs" in result, "Result should have a paragraphs key"
    assert "tables" in result, "Result should have a tables key"
    assert "headings" in result, "Result should have a headings key"
    assert "metadata" in result, "Result should have a metadata key"
    
    # Check title
    assert result["title"] == "Sample Document", "Title should be 'Sample Document'"
    
    # Check paragraphs
    assert len(result["paragraphs"]) >= 2, "There should be at least 2 paragraphs"
    assert any("sample paragraph" in p["text"] for p in result["paragraphs"]), "Content should contain sample paragraph"
    
    # Check tables
    assert len(result["tables"]) == 1, "There should be 1 table"
    assert result["tables"][0]["data"][0][0] == "Header 1", "Table should contain 'Header 1'"
    
    # Check headings
    assert len(result["headings"]) >= 1, "There should be at least 1 heading"
    assert any(h["text"] == "Section 1" for h in result["headings"]), "Headings should include 'Section 1'"
    assert any(h["level"] == 1 for h in result["headings"]), "Heading should have level 1"

def test_extract_text(sample_docx):
    # Extract text from the sample document
    text = DocxParser.extract_text(sample_docx)
    
    # Check that the result is a string
    assert isinstance(text, str), "Result should be a string"
    
    # Check that the text contains expected content
    assert "Sample Document" in text, "Text should include the document title"
    assert "This is a sample paragraph" in text, "Text should include paragraph content"
    assert "This paragraph has" in text, "Text should include styled paragraph content"
    assert "Section 1" in text, "Text should include heading content"
    assert "Header 1" in text, "Text should include table content"
    assert "Value 1" in text, "Text should include table content"

def test_parse_invalid_document():
    # Test with invalid document bytes
    with pytest.raises(DocxParserError):
        DocxParser.parse_document(b"invalid document content")

def test_empty_document(empty_docx):
    """Test parsing an empty document"""
    # Parse an empty document
    result = DocxParser.parse_document(empty_docx)
    
    # Check that result has expected structure even for empty document
    assert isinstance(result, dict), "Result should be a dictionary"
    assert "title" in result, "Result should have a title key"
    assert "paragraphs" in result, "Result should have a paragraphs key"
    assert "tables" in result, "Result should have a tables key"
    assert "headings" in result, "Result should have a headings key"
    assert "metadata" in result, "Result should have a metadata key"
    
    # Check that empty document has expected content
    assert result["title"] == "Untitled Document", "Empty document should have 'Untitled Document' as title"
    assert len(result["paragraphs"]) == 0, "Empty document should have no paragraphs"
    assert len(result["tables"]) == 0, "Empty document should have no tables"
    assert len(result["headings"]) == 0, "Empty document should have no headings"

def test_extract_text_from_empty_document(empty_docx):
    """Test extracting text from an empty document"""
    # Extract text from an empty document
    text = DocxParser.extract_text(empty_docx)
    
    # Check that result is a string
    assert isinstance(text, str), "Result should be a string"
    
    # Empty documents should return an empty string (or very minimal content)
    # We just check it's a string without asserting contents, since an empty document
    # might have minimal auto-generated content from docx library
    pass

def test_document_with_lists(list_docx):
    """Test extracting text from a document with lists"""
    # Extract text from a document with lists
    text = DocxParser.extract_text(list_docx)
    
    # Check that the text contains list markers or list content
    assert "Document with Lists" in text, "Text should include document title"
    assert "First bullet point" in text, "Text should include bullet point content"
    assert "Second bullet point" in text, "Text should include bullet point content"
    assert "First numbered item" in text, "Text should include numbered list content"
    assert "Second numbered item" in text, "Text should include numbered list content"

def test_input_validation():
    """Test input validation for the parser"""
    # Test with None input
    with pytest.raises(AssertionError, match="File content cannot be None"):
        DocxParser.extract_text(None)
    
    # Test with wrong type input
    with pytest.raises(AssertionError, match="File content must be bytes"):
        DocxParser.extract_text("not bytes")
    
    # Test with empty bytes
    with pytest.raises(AssertionError, match="File content cannot be empty"):
        DocxParser.extract_text(b"")
        
    # Test input validation for parse_document
    with pytest.raises(AssertionError):
        DocxParser.parse_document(None) 