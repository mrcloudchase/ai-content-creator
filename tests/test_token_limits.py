import pytest
from fastapi.testclient import TestClient
import io
import os
from unittest.mock import patch, MagicMock
from app.main import app
from docx import Document
from app.services.token_service import TokenService
from app.services.docx_parser import DocxParser, TokenLimitError

client = TestClient(app)

def create_large_document(repeat_count=5000):
    """
    Create a test document that would exceed token limits
    """
    doc = Document()
    
    # Add a title
    doc.add_heading('Very Large Test Document', 0)
    
    # Add a lot of paragraphs
    for i in range(repeat_count):
        doc.add_paragraph(f"This is paragraph {i+1} with some text that will contribute to the token count. " * 5)
    
    # Save the document to a bytesIO object
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream

# Store the original extract_text method
original_extract_text = DocxParser.extract_text

@pytest.mark.parametrize("token_limit_exceeded", [True, False])
def test_document_parser_token_limit(token_limit_exceeded):
    """Test that documents exceeding token limits return appropriate errors"""
    try:
        # Replace the extract_text method with our mock
        def mock_extract_text(file_content):
            if token_limit_exceeded:
                raise TokenLimitError("Document exceeds token limit: 5000 tokens (limit: 1000). Use the token counting endpoint for more information.")
            else:
                return "This is the extracted text."
        
        DocxParser.extract_text = mock_extract_text
        
        # Create a simple document
        doc = Document()
        doc.add_paragraph("Test document")
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        
        # Make the request
        response = client.post(
            "/api/v1/documents/extract-text",
            files={"file": ("test.docx", file_stream, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
        )
        
        # Check the response
        expected_status = 413 if token_limit_exceeded else 200
        assert response.status_code == expected_status
        
        # If we expect an error, verify it's the token limit error
        if token_limit_exceeded:
            json_response = response.json()
            assert "token_limit_exceeded" in json_response["error_type"]
            assert "Document exceeds token limit" in json_response["detail"]
        else:
            # For successful response, check for document field
            json_response = response.json()
            assert "document" in json_response
            assert json_response["document"] == "This is the extracted text."
    finally:
        # Restore the original method
        DocxParser.extract_text = original_extract_text

@pytest.mark.xfail(reason="This test requires a real large document and may be slow")
def test_very_large_document_real_tokens():
    """Test a real large document without mocks to verify actual token counting"""
    # Create a large document
    file_stream = create_large_document()
    
    # Make the request
    response = client.post(
        "/api/v1/documents/extract-text",
        files={"file": ("large_doc.docx", file_stream, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    
    # Check the response - we expect 413 because the document should be large enough
    # to exceed the default max_tokens (usually 1000)
    assert response.status_code == 413
    json_response = response.json()
    assert "token_limit_exceeded" in json_response["error_type"] 