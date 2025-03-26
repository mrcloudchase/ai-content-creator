import pytest
from fastapi.testclient import TestClient
from app.main import app
import io
import docx
import os

client = TestClient(app)

@pytest.fixture
def sample_docx_file():
    """Create a sample .docx file for testing"""
    doc = docx.Document()
    
    # Add title
    doc.add_heading('API Test Document', 0)
    
    # Add a paragraph
    doc.add_paragraph('This is a test paragraph for the API endpoint.')
    
    # Add a heading
    doc.add_heading('Test Section', level=1)
    
    # Add a table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = 'Header 1'
    table.cell(0, 1).text = 'Header 2'
    table.cell(1, 0).text = 'Value 1'
    table.cell(1, 1).text = 'Value 2'
    
    # Save to bytes
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    return docx_bytes

@pytest.fixture
def empty_docx_file():
    """Create an empty .docx file for testing"""
    doc = docx.Document()
    
    # Save to bytes
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    return docx_bytes

def test_root_endpoint():
    """Test the root endpoint"""
    response = client.get("/")
    
    assert response.status_code == 200
    data = response.json()
    assert "name" in data
    assert "version" in data
    assert "documentation" in data
    assert "endpoints" in data

def test_health_endpoint():
    """Test the health check endpoint"""
    response = client.get("/health")
    
    assert response.status_code == 200
    data = response.json()
    assert data["status"] == "healthy"
    assert "version" in data

def test_extract_text_endpoint(sample_docx_file):
    """Test the document text extraction endpoint"""
    response = client.post(
        "/api/v1/documents/extract-text",
        files={"file": ("test.docx", sample_docx_file, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    
    # Check status code
    assert response.status_code == 200, f"Expected status code 200, got {response.status_code}"
    
    # Check that the response is plain text
    assert "API Test Document" in response.text, "Text should include the document title"
    assert "This is a test paragraph" in response.text, "Text should include paragraph content"
    assert "Test Section" in response.text, "Text should include section headings"
    assert "Header 1" in response.text, "Text should include table content"
    assert "Value 1" in response.text, "Text should include table content"

def test_invalid_file_endpoint():
    """Test the document endpoint with an invalid file type"""
    # Create a text file
    text_file = io.BytesIO(b"This is not a docx file")
    
    response = client.post(
        "/api/v1/documents/extract-text",
        files={"file": ("test.txt", text_file, "text/plain")}
    )
    
    # Check status code
    assert response.status_code == 400, f"Expected status code 400, got {response.status_code}"
    
    # Check error message
    data = response.json()
    assert "detail" in data, "Error response should have a detail key"
    assert "Invalid file format" in data["detail"], "Error should mention invalid file format"

def test_empty_docx_file(empty_docx_file):
    """Test the document endpoint with an empty docx file"""
    response = client.post(
        "/api/v1/documents/extract-text",
        files={"file": ("empty.docx", empty_docx_file, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    
    # Check status code - should succeed with empty string
    assert response.status_code == 200, f"Expected status code 200, got {response.status_code}"
    
    # Empty doc should return empty text (or possibly minimal content)
    # Just check response is successful without asserting exact content
    assert isinstance(response.text, str), "Response should be a string"

def test_missing_file():
    """Test the document endpoint without providing a file"""
    response = client.post("/api/v1/documents/extract-text")
    
    # Check status code
    assert response.status_code == 422, f"Expected status code 422, got {response.status_code}"
    
    # Check error message
    data = response.json()
    assert "detail" in data, "Error response should have a detail key"
    assert "field required" in str(data["detail"]).lower(), "Error should mention required field"

def test_incorrect_content_type():
    """Test the document endpoint with correct extension but incorrect content type"""
    # Create a fake docx file that's actually a text file
    fake_docx = io.BytesIO(b"This is not really a docx file")
    
    response = client.post(
        "/api/v1/documents/extract-text",
        files={"file": ("fake.docx", fake_docx, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    
    # Should fail with validation error when trying to parse it
    assert response.status_code == 422, f"Expected status code 422, got {response.status_code}"
    
    # Check error message
    data = response.json()
    assert "detail" in data, "Error response should have a detail key"
    assert "error parsing document" in data["detail"].lower(), "Error should mention parsing error"

def test_empty_file_content():
    """Test the document endpoint with an empty file content"""
    empty_file = io.BytesIO(b"")
    
    response = client.post(
        "/api/v1/documents/extract-text",
        files={"file": ("empty.docx", empty_file, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    
    # Check status code
    assert response.status_code == 400, f"Expected status code 400, got {response.status_code}"
    
    # Check error message
    data = response.json()
    assert "detail" in data, "Error response should have a detail key"
    assert "empty file" in data["detail"].lower(), "Error should mention empty file" 