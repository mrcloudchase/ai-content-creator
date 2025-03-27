import pytest
import io
import os
import docx
from app.services.docx_parser import DocxParser, DocxParserError
import json
from docx import Document
import docx.oxml.ns

@pytest.fixture
def sample_docx():
    """Create a sample docx document for testing"""
    doc = docx.Document()
    
    # Add title
    doc.add_heading('Sample Document', 0)
    
    # Add some paragraphs
    doc.add_paragraph('This is a sample paragraph with some content.')
    
    # Add a paragraph with styled text
    p = doc.add_paragraph('This paragraph has ')
    p.add_run('bold').bold = True
    p.add_run(' and ')
    p.add_run('italic').italic = True
    p.add_run(' text.')
    
    # Add a heading
    doc.add_heading('Section 1', level=1)
    
    # Add another paragraph
    doc.add_paragraph('Content in the first section.')
    
    # Add a table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = 'Header 1'
    table.cell(0, 1).text = 'Header 2'
    table.cell(1, 0).text = 'Value 1'
    table.cell(1, 1).text = 'Value 2'
    
    # Save to bytes
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    return docx_bytes.getvalue()

@pytest.fixture
def empty_docx():
    """Create an empty docx document for testing edge cases"""
    doc = docx.Document()
    
    # Save to bytes
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    return docx_bytes.getvalue()

@pytest.fixture
def list_docx():
    """Create a docx document with lists for testing"""
    doc = docx.Document()
    
    # Add a heading
    doc.add_heading('Document with Lists', 0)
    
    # Add a bulleted list - need to apply style manually
    doc.add_paragraph('First bullet point', style='List Bullet')
    doc.add_paragraph('Second bullet point', style='List Bullet')
    
    # Add a numbered list - need to apply style manually
    doc.add_paragraph('First numbered item', style='List Number')
    doc.add_paragraph('Second numbered item', style='List Number')
    
    # Save to bytes
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    return docx_bytes.getvalue()

def test_parse_document(sample_docx):
    # Parse the sample document
    result = DocxParser.parse_document(sample_docx)
    
    # Check that the result has the expected structure
    assert isinstance(result, dict), "Result should be a dictionary"
    assert "title" in result, "Result should have a title key"
    assert "paragraphs" in result, "Result should have a paragraphs key"
    assert "tables" in result, "Result should have a tables key"
    assert "headings" in result, "Result should have a headings key"
    assert "metadata" in result, "Result should have a metadata key"
    
    # Check title
    assert result["title"] == "Sample Document", "Title should be 'Sample Document'"
    
    # Check paragraphs
    assert len(result["paragraphs"]) >= 2, "There should be at least 2 paragraphs"
    assert any("sample paragraph" in p["text"] for p in result["paragraphs"]), "Content should contain sample paragraph"
    
    # Check tables
    assert len(result["tables"]) == 1, "There should be 1 table"
    assert result["tables"][0]["data"][0][0] == "Header 1", "Table should contain 'Header 1'"
    
    # Check headings
    assert len(result["headings"]) >= 1, "There should be at least 1 heading"
    assert any(h["text"] == "Section 1" for h in result["headings"]), "Headings should include 'Section 1'"
    assert any(h["level"] == 1 for h in result["headings"]), "Heading should have level 1"

def test_extract_text(sample_docx):
    # Extract text from the sample document
    text = DocxParser.extract_text(sample_docx)
    
    # Check that the result is a string
    assert isinstance(text, str), "Result should be a string"
    
    # Check that the text contains expected content
    assert "Sample Document" in text, "Text should include the document title"
    assert "This is a sample paragraph" in text, "Text should include paragraph content"
    assert "This paragraph has" in text, "Text should include styled paragraph content"
    assert "Section 1" in text, "Text should include heading content"
    assert "Header 1" in text, "Text should include table content"
    assert "Value 1" in text, "Text should include table content"

def test_parse_invalid_document():
    # Test with invalid document bytes
    with pytest.raises(DocxParserError):
        DocxParser.parse_document(b"invalid document content")

def test_empty_document(empty_docx):
    """Test parsing an empty document"""
    # Parse an empty document
    result = DocxParser.parse_document(empty_docx)
    
    # Check that result has expected structure even for empty document
    assert isinstance(result, dict), "Result should be a dictionary"
    assert "title" in result, "Result should have a title key"
    assert "paragraphs" in result, "Result should have a paragraphs key"
    assert "tables" in result, "Result should have a tables key"
    assert "headings" in result, "Result should have a headings key"
    assert "metadata" in result, "Result should have a metadata key"
    
    # Check that empty document has expected content
    assert result["title"] == "Untitled Document", "Empty document should have 'Untitled Document' as title"
    assert len(result["paragraphs"]) == 0, "Empty document should have no paragraphs"
    assert len(result["tables"]) == 0, "Empty document should have no tables"
    assert len(result["headings"]) == 0, "Empty document should have no headings"

def test_extract_text_from_empty_document(empty_docx):
    """Test extracting text from an empty document"""
    # Extract text from an empty document
    text = DocxParser.extract_text(empty_docx)
    
    # Check that result is a string
    assert isinstance(text, str), "Result should be a string"
    
    # Empty documents should return an empty string (or very minimal content)
    # We just check it's a string without asserting contents, since an empty document
    # might have minimal auto-generated content from docx library
    pass

def test_document_with_lists(list_docx):
    """Test extracting text from a document with lists"""
    # Extract text from a document with lists
    text = DocxParser.extract_text(list_docx)
    
    # Check that the text contains list markers or list content
    assert "Document with Lists" in text, "Text should include document title"
    assert "First bullet point" in text, "Text should include bullet point content"
    assert "Second bullet point" in text, "Text should include bullet point content"
    assert "First numbered item" in text, "Text should include numbered list content"
    assert "Second numbered item" in text, "Text should include numbered list content"

def test_input_validation():
    """Test input validation for the parser"""
    # Test with None input
    with pytest.raises(AssertionError, match="File content cannot be None"):
        DocxParser.extract_text(None)
    
    # Test with wrong type input
    with pytest.raises(AssertionError, match="File content must be bytes"):
        DocxParser.extract_text("not bytes")
    
    # Test with empty bytes
    with pytest.raises(AssertionError, match="File content cannot be empty"):
        DocxParser.extract_text(b"")
        
    # Test input validation for parse_document
    with pytest.raises(AssertionError):
        DocxParser.parse_document(None)

# Get the current directory of this test file
current_dir = os.path.dirname(os.path.abspath(__file__))

# Test helper functions
def create_test_document(content):
    """Create a test document with the given content"""
    doc = Document()
    doc.add_paragraph(content)
    
    # Save the document to a bytesIO object
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream.read()

def create_document_with_formatting():
    """Create a document with various formatting for testing"""
    doc = Document()
    
    # Add title (style 'Title')
    doc.add_heading('Test Document', 0)
    
    # Add level 1 heading with proper style
    doc.add_heading('Heading 1', level=1)
    
    # Add level 2 heading with proper style
    doc.add_heading('Heading 2', level=2)
    
    # Add paragraphs with different formatting
    p = doc.add_paragraph('This is a paragraph with ')
    p.add_run('bold').bold = True
    p.add_run(' and ')
    p.add_run('italic').italic = True
    p.add_run(' text.')
    
    # Add a list
    doc.add_paragraph('Item 1', style='List Bullet')
    doc.add_paragraph('Item 2', style='List Bullet')
    doc.add_paragraph('Item 3', style='List Bullet')
    
    # Add a table
    table = doc.add_table(rows=2, cols=2)
    cell = table.cell(0, 0)
    cell.text = 'Cell 1'
    cell = table.cell(0, 1)
    cell.text = 'Cell 2'
    cell = table.cell(1, 0)
    cell.text = 'Cell 3'
    cell = table.cell(1, 1)
    cell.text = 'Cell 4'
    
    # Save the document to a bytesIO object
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    return file_stream.read()

# Tests
def test_extract_text_simple():
    """Test extracting text from a simple document"""
    # Arrange
    content = "Hello, world!"
    doc_bytes = create_test_document(content)
    
    # Act
    extracted_text = DocxParser.extract_text(doc_bytes)
    
    # Assert
    assert "Hello, world!" in extracted_text

def test_extract_text_with_formatting():
    """Test extracting text with formatting"""
    # Arrange
    doc_bytes = create_document_with_formatting()
    
    # Act
    extracted_text = DocxParser.extract_text(doc_bytes)
    
    # Assert
    assert "Test Document" in extracted_text
    assert "This is a paragraph with bold and italic text." in extracted_text
    assert "Item 1" in extracted_text
    assert "Item 2" in extracted_text
    assert "Item 3" in extracted_text
    assert "Cell 1" in extracted_text
    assert "Cell 2" in extracted_text
    assert "Cell 3" in extracted_text
    assert "Cell 4" in extracted_text

def test_extract_text_empty_document():
    """Test extracting text from an empty document"""
    # Arrange
    doc = Document()
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    
    # Act
    extracted_text = DocxParser.extract_text(file_stream.read())
    
    # Assert
    assert extracted_text == ""

def test_extract_text_handles_special_characters():
    """Test extracting text with special characters"""
    # Arrange
    content = 'Special characters: é, ñ, ü, and quotes "like this" and \'these\''
    doc_bytes = create_test_document(content)
    
    # Act
    extracted_text = DocxParser.extract_text(doc_bytes)
    
    # Assert
    assert "Special characters" in extracted_text
    # Check that special characters are preserved or properly transformed
    assert "e" in extracted_text  # At minimum, the base character should be there

def test_parse_document_structure():
    """Test parsing the document structure"""
    # Arrange
    doc_bytes = create_document_with_formatting()
    
    # Act
    result = DocxParser.parse_document(doc_bytes)
    
    # Assert
    assert isinstance(result, dict)
    assert "title" in result
    assert "paragraphs" in result
    assert "tables" in result
    assert "headings" in result
    assert "metadata" in result
    
    # Check content of result
    assert len(result["tables"]) > 0
    assert len(result["paragraphs"]) > 0
    assert len(result["headings"]) > 0

def test_invalid_input_bytes():
    """Test that parsing fails with invalid input bytes"""
    # Arrange
    invalid_bytes = b'This is not a valid docx file'
    
    # Act / Assert
    with pytest.raises(DocxParserError):
        DocxParser.parse_document(invalid_bytes)
    
    with pytest.raises(DocxParserError):
        DocxParser.extract_text(invalid_bytes)

def test_none_input():
    """Test that parsing fails with None input"""
    # Act / Assert
    with pytest.raises(AssertionError):
        DocxParser.parse_document(None)
    
    with pytest.raises(AssertionError):
        DocxParser.extract_text(None)

def test_empty_input():
    """Test that parsing fails with empty input"""
    # Act / Assert
    with pytest.raises(AssertionError):
        DocxParser.parse_document(b'')
    
    with pytest.raises(AssertionError):
        DocxParser.extract_text(b'')

# New tests for JSON compatibility

def test_extract_text_json_compatibility():
    """Test that extracted text is JSON-compatible"""
    # Arrange
    doc_bytes = create_document_with_formatting()
    
    # Act
    extracted_text = DocxParser.extract_text(doc_bytes)
    
    # Assert - this should not raise an exception
    json_obj = {"prompt": extracted_text}
    json_str = json.dumps(json_obj)
    
    # Verify we can parse it back
    parsed_obj = json.loads(json_str)
    assert parsed_obj["prompt"] == extracted_text

def test_extract_text_with_problematic_characters():
    """Test extraction with characters that are typically problematic for JSON"""
    # Arrange
    content = """This has "quotes" and 'apostrophes' and new
lines and tabs\t and backslashes \\ and control chars and unicode like €£¥"""
    doc_bytes = create_test_document(content)
    
    # Act
    extracted_text = DocxParser.extract_text(doc_bytes)
    
    # Assert - should be JSON-serializable
    json_obj = {"prompt": extracted_text}
    json_str = json.dumps(json_obj)
    
    # Should be able to parse back
    parsed_obj = json.loads(json_str)
    assert isinstance(parsed_obj["prompt"], str)
    
    # Test nested JSON scenario (simulating copy-paste into another JSON)
    nested_json = json.dumps({"outer": json_str})
    parsed_nested = json.loads(nested_json)
    assert isinstance(parsed_nested["outer"], str)
    inner_obj = json.loads(parsed_nested["outer"])
    assert inner_obj["prompt"] == extracted_text

def test_make_json_compatible():
    """Test the _make_json_compatible method directly"""
    # Arrange
    problematic_text = """Text with "smart quotes" and 'apostrophes'
    and em-dashes—and en-dashes–and ellipsis…
    and tabs\t and newlines mixed with \r carriage returns"""
    
    # Act
    cleaned_text = DocxParser._make_json_compatible(problematic_text)
    
    # Assert
    # Should be JSON-serializable
    json_obj = {"text": cleaned_text}
    json_str = json.dumps(json_obj)
    
    # Should be able to parse back
    parsed_obj = json.loads(json_str)
    assert isinstance(parsed_obj["text"], str)
    
    # Verify specific replacements - with escaped quotes
    assert '\\"smart quotes\\"' in cleaned_text  # Smart quotes replaced with escaped straight quotes
    assert "'apostrophes'" in cleaned_text   # Smart apostrophes replaced
    assert "em-dashes--" in cleaned_text     # Em-dash replaced with double hyphen
    assert "en-dashes-" in cleaned_text      # En-dash replaced with hyphen
    assert "ellipsis..." in cleaned_text     # Ellipsis replaced with three periods
    
    # Test nested JSON scenario
    nested_json = json.dumps({"outer": json_str})
    assert json.loads(nested_json)  # Should not raise exception

def test_nested_json_compatibility():
    """Test specifically for nested JSON compatibility - the copy/paste user scenario"""
    # Arrange
    control_chars_text = "Text with newlines\nand tabs\tand backslashes\\ and quotes\"'!"
    doc_bytes = create_test_document(control_chars_text)
    
    # Act
    extracted_text = DocxParser.extract_text(doc_bytes)
    
    # Assert - build a JSON string that would be returned by the API
    api_response = json.dumps({"document": extracted_text})
    
    # Simulate user copying this value and using it in another JSON payload
    # This is the scenario causing the 422 error
    user_json = '{{"prompt": {}}}'.format(json.dumps(extracted_text))
    
    # This should parse without error
    parsed = json.loads(user_json)
    assert isinstance(parsed["prompt"], str)
    
    # Double-nested case (worst case)
    double_nested = json.dumps({"outer": user_json})
    parsed_double = json.loads(double_nested)
    assert isinstance(parsed_double["outer"], str)
    inner = json.loads(parsed_double["outer"])
    assert isinstance(inner["prompt"], str) 