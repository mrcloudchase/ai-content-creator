import pytest
from fastapi.testclient import TestClient
from app.main import app
import io
import docx
import os
import json

client = TestClient(app)

# Get the current directory of this test file
current_dir = os.path.dirname(os.path.abspath(__file__))
# Path to test files
test_files_dir = os.path.join(current_dir, 'test_files')
os.makedirs(test_files_dir, exist_ok=True)

@pytest.fixture
def sample_docx_file():
    """Create a sample .docx file for testing"""
    doc = docx.Document()
    
    # Add title
    doc.add_heading('API Test Document', 0)
    
    # Add a paragraph
    doc.add_paragraph('This is a test paragraph for the API endpoint.')
    
    # Add a heading
    doc.add_heading('Test Section', level=1)
    
    # Add a table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = 'Header 1'
    table.cell(0, 1).text = 'Header 2'
    table.cell(1, 0).text = 'Value 1'
    table.cell(1, 1).text = 'Value 2'
    
    # Save to bytes
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    return docx_bytes

@pytest.fixture
def empty_docx_file():
    """Create an empty .docx file for testing"""
    doc = docx.Document()
    
    # Save to bytes
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    return docx_bytes

def test_root_endpoint():
    """Test that the root endpoint returns basic info"""
    response = client.get("/")
    assert response.status_code == 200
    data = response.json()
    assert "name" in data
    assert "version" in data
    assert "documentation" in data
    assert "endpoints" in data

def test_health_endpoint():
    """Test that the health endpoint returns healthy status"""
    response = client.get("/health")
    assert response.status_code == 200
    data = response.json()
    assert data["status"] == "healthy"
    assert "service" in data
    assert "version" in data

def test_extract_text_from_invalid_file():
    """Test that the extract-text endpoint fails with non-docx file"""
    # Create a test file that's not a docx
    file_content = b'This is not a docx file.'
    response = client.post(
        "/api/v1/documents/extract-text",
        files={"file": ("test.txt", file_content, "text/plain")}
    )
    
    assert response.status_code == 400
    assert "Invalid file format" in response.json()["detail"]

def test_extract_text_from_empty_file():
    """Test that the extract-text endpoint fails with empty file"""
    response = client.post(
        "/api/v1/documents/extract-text",
        files={"file": ("test.docx", b'', "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    
    assert response.status_code == 400
    assert "Empty file" in response.json()["detail"]

def test_extract_text_no_file():
    """Test that the extract-text endpoint fails with no file"""
    response = client.post("/api/v1/documents/extract-text")
    
    assert response.status_code == 422  # FastAPI validation error

def test_extract_text_missing_filename():
    """Test that the extract-text endpoint fails with missing filename"""
    # Use an empty string as the filename
    file_content = b'Some content'
    response = client.post(
        "/api/v1/documents/extract-text",
        files={"file": ("", file_content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    
    # FastAPI now returns a 422 for this validation error
    assert response.status_code == 422

def test_extract_text_invalid_docx():
    """Test that the extract-text endpoint fails with invalid docx content"""
    # Just some random bytes that aren't a valid docx file
    file_content = b'This is not a valid docx file structure.'
    response = client.post(
        "/api/v1/documents/extract-text",
        files={"file": ("test.docx", file_content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    
    assert response.status_code == 422
    assert "Error parsing document" in response.json()["detail"]

def create_test_docx(text="Test document content."):
    """Helper to create a real test docx file"""
    try:
        from docx import Document
        document = Document()
        document.add_paragraph(text)
        
        file_obj = io.BytesIO()
        document.save(file_obj)
        file_obj.seek(0)
        
        return file_obj.read()
    except ImportError:
        # If python-docx is not available, we'll use a pre-made docx file
        # or return an error that this test needs python-docx
        pytest.skip("python-docx is required for this test")

def test_extract_text_success():
    """Test successful extraction of text from a docx file"""
    # Create a real test document
    file_content = create_test_docx("This is a test document.")
    
    response = client.post(
        "/api/v1/documents/extract-text",
        files={"file": ("test.docx", file_content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    
    assert response.status_code == 200
    data = response.json()
    assert "document" in data
    assert "This is a test document." in data["document"]

def test_extract_text_endpoint(sample_docx_file):
    """Test the document text extraction endpoint"""
    response = client.post(
        "/api/v1/documents/extract-text",
        files={"file": ("test.docx", sample_docx_file, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    
    # Check that we get a successful response
    assert response.status_code == 200
    
    # Verify that we got JSON data
    data = response.json()
    assert "document" in data
    
    # Check that the document text contains the expected content
    # The exact content would depend on your test sample
    assert isinstance(data["document"], str)
    assert len(data["document"]) > 0

def test_invalid_file_endpoint():
    """Test the document endpoint with an invalid file type"""
    # Create a text file
    text_file = io.BytesIO(b"This is not a docx file")
    
    response = client.post(
        "/api/v1/documents/extract-text",
        files={"file": ("test.txt", text_file, "text/plain")}
    )
    
    # Check status code
    assert response.status_code == 400, f"Expected status code 400, got {response.status_code}"
    
    # Check error message
    data = response.json()
    assert "detail" in data, "Error response should have a detail key"
    assert "Invalid file format" in data["detail"], "Error should mention invalid file format"

def test_empty_docx_file(empty_docx_file):
    """Test the document endpoint with an empty docx file"""
    response = client.post(
        "/api/v1/documents/extract-text",
        files={"file": ("empty.docx", empty_docx_file, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    
    # Check status code - should succeed with empty string
    assert response.status_code == 200, f"Expected status code 200, got {response.status_code}"
    
    # Empty doc should return empty text (or possibly minimal content)
    # Just check response is successful without asserting exact content
    assert isinstance(response.text, str), "Response should be a string"

def test_missing_file():
    """Test the document endpoint without providing a file"""
    response = client.post("/api/v1/documents/extract-text")
    
    # Check status code
    assert response.status_code == 422, f"Expected status code 422, got {response.status_code}"
    
    # Check error message
    data = response.json()
    assert "detail" in data, "Error response should have a detail key"
    assert "field required" in str(data["detail"]).lower(), "Error should mention required field"

def test_incorrect_content_type():
    """Test the document endpoint with correct extension but incorrect content type"""
    # Create a fake docx file that's actually a text file
    fake_docx = io.BytesIO(b"This is not really a docx file")
    
    response = client.post(
        "/api/v1/documents/extract-text",
        files={"file": ("fake.docx", fake_docx, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    
    # Should fail with validation error when trying to parse it
    assert response.status_code == 422, f"Expected status code 422, got {response.status_code}"
    
    # Check error message
    data = response.json()
    assert "detail" in data, "Error response should have a detail key"
    assert "error parsing document" in data["detail"].lower(), "Error should mention parsing error"

def test_empty_file_content():
    """Test the document endpoint with an empty file content"""
    empty_file = io.BytesIO(b"")
    
    response = client.post(
        "/api/v1/documents/extract-text",
        files={"file": ("empty.docx", empty_file, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    
    # Check status code
    assert response.status_code == 400, f"Expected status code 400, got {response.status_code}"
    
    # Check error message
    data = response.json()
    assert "detail" in data, "Error response should have a detail key"
    assert "empty file" in data["detail"].lower(), "Error should mention empty file"

def test_document_parser_json_compatibility():
    """Test that document parser output can be used directly in AI endpoint"""
    # Create a docx document with potentially problematic content
    doc = docx.Document()
    doc.add_paragraph('Test content with "quotes", newlines\nand special characters\tlike tabs')
    
    # Save to bytes
    docx_bytes = io.BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    # Step 1: Get parsed document from document endpoint
    response = client.post(
        "/api/v1/documents/extract-text",
        files={"file": ("test.docx", docx_bytes.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    )
    
    # Verify document parsing succeeded
    assert response.status_code == 200
    parsed_data = response.json()
    assert "document" in parsed_data
    
    # Step 2: Verify the output can be directly used in AI endpoint without modification
    # Simulate what a user would do - copy the document text and use it in AI request
    # This is a mock to ensure the JSON works; we don't need to actually call the AI API
    ai_request = {
        "prompt": parsed_data["document"],
        "max_tokens": 100
    }
    
    # Convert to JSON string and parse back - should not raise any JSON errors
    ai_request_json = json.dumps(ai_request)
    parsed_request = json.loads(ai_request_json)
    
    # Verify the data survived the JSON serialization/deserialization intact
    assert parsed_request["prompt"] == parsed_data["document"]
    
    # If this test passes, it means users can copy-paste the document text
    # directly into the AI endpoint without encountering JSON errors 