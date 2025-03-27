#!/usr/bin/env python3
import requests
import docx
import io
import json

def create_full_prd_document():
    """Create a detailed PRD document from the full text example"""
    doc = docx.Document()
    
    # Full PRD text
    prd_text = """
Azure Kubernetes Service (AKS) PRD
Feature: Support for CiliumEndpointSlices with Azure CNI by Cilium
Document Version: 1.0
Document Date: April 1, 2025
Prepared By: AKS Product Team

1. Document Overview

This PRD defines the requirements for integrating CiliumEndpointSlices with Azure CNI by Cilium into Azure Kubernetes Service (AKS). The goal is to enhance AKS's networking capabilities by leveraging Cilium's dynamic endpoint management and advanced network policy enforcement. This document details the business rationale, feature description, user scenarios, functional and non-functional requirements, assumptions, dependencies, timeline, and future considerations to guide design, development, testing, and release.

2. Business Objectives

• Enhanced Networking Efficiency:

  • Enable dynamic grouping of service endpoints for rapid service discovery and efficient resource utilization.

• Improved Security & Policy Enforcement:

  • Utilize Cilium's native network policy engine to enforce granular security rules that align with enterprise compliance requirements.

• Optimized Developer & Operator Experience:

  • Provide a seamless configuration experience via Azure Portal, CLI, and APIs that reduces operational overhead and accelerates deployments.

• Competitive Differentiation:

  • Position AKS as a leader in managed Kubernetes by integrating state-of-the-art networking features that meet the evolving needs of cloud-native applications.

• Operational Cost Efficiency:

  • Enhance the scalability and stability of clusters to reduce downtime and lower management costs, particularly in large-scale deployments.

3. Feature Overview

The Support for CiliumEndpointSlices with Azure CNI by Cilium feature will enable AKS to:

• Dynamically Manage Endpoints:
 Automatically group and manage pod endpoints into slices, which allows for faster lookup, improved scalability, and efficient load balancing.

• Enforce Advanced Network Policies:
 Leverage Cilium's rich network policy language to implement fine-grained security policies across dynamic endpoint slices.

• Enhance Observability:
 Provide detailed logging, metrics, and diagnostic information through integrated Azure Monitor and Cilium observability tools, enabling proactive troubleshooting and performance tuning.

• Offer Configurability:
 Allow customers to enable or disable the feature during cluster provisioning, with customizable parameters (e.g., refresh intervals, slice size thresholds) to suit diverse operational needs.
    """
    
    # Split the text by lines and add each line as a paragraph
    for line in prd_text.strip().split('\n'):
        # Skip empty lines
        if not line.strip():
            continue
            
        # Check if it's a heading
        if line.startswith('Azure Kubernetes Service'):
            doc.add_heading(line, 0)
        elif line.startswith(('1.', '2.', '3.')):
            doc.add_heading(line, level=1)
        elif line.strip().startswith('•'):
            doc.add_paragraph(line, style='List Bullet')
        else:
            doc.add_paragraph(line)
    
    # Save to a BytesIO object
    f = io.BytesIO()
    doc.save(f)
    f.seek(0)
    
    return f.read()

def main():
    """Test the document parser with the PRD document and then use it with the customer intent endpoint"""
    print("Creating PRD document...")
    doc_bytes = create_full_prd_document()
    
    # Step 1: Parse the document
    print("\nStep 1: Sending PRD document for parsing...")
    files = {"file": ("prd.docx", doc_bytes, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")}
    parse_response = requests.post("http://localhost:8000/api/v1/documents/extract-text", files=files)
    
    if parse_response.status_code != 200:
        print(f"Error parsing document: {parse_response.status_code}")
        print(parse_response.text)
        return
    
    print(f"Document parsed successfully: {parse_response.status_code}")
    parsed_data = parse_response.json()
    
    # Print a preview of the document text
    preview = parsed_data["document"][:200] + "..." if len(parsed_data["document"]) > 200 else parsed_data["document"]
    print(f"\nDocument preview:\n{'-'*50}\n{preview}\n{'-'*50}")
    
    # Step 2: Test with AI endpoint
    print("\nStep 2: Testing with AI endpoint...")
    ai_request = {
        "prompt": "Summarize the following PRD document in 3-4 bullet points:\n\n" + parsed_data["document"],
        "max_tokens": 300,
        "temperature": 0.5
    }
    
    # Send to AI endpoint
    ai_response = requests.post("http://localhost:8000/api/v1/ai/completions", json=ai_request)
    
    if ai_response.status_code != 200:
        print(f"❌ AI endpoint call failed: {ai_response.status_code}")
        print(ai_response.text)
    else:
        print(f"✅ AI endpoint call successful: {ai_response.status_code}")
        ai_data = ai_response.json()
        
        print("\nAI Response:")
        print(f"{'-'*50}\n{ai_data['text']}\n{'-'*50}")
    
    # Step 3: Test with Customer Intent endpoint
    print("\nStep 3: Generating customer intent from the document...")
    intent_request = {
        "document_text": parsed_data["document"],
        "max_tokens": 150,
        "temperature": 0.7
    }
    
    intent_response = requests.post("http://localhost:8000/api/v1/customer-intent/generate", json=intent_request)
    
    if intent_response.status_code != 200:
        print(f"❌ Customer intent generation failed: {intent_response.status_code}")
        print(intent_response.text)
    else:
        print(f"✅ Customer intent generated successfully: {intent_response.status_code}")
        intent_data = intent_response.json()
        
        print("\nCustomer Intent:")
        print(f"{'-'*50}\n{intent_data['intent']}\n{'-'*50}")
    
    # Step 4: Generate intent for specific user types
    user_types = ["cluster operator", "network administrator", "security engineer"]
    
    for user_type in user_types:
        print(f"\nGenerating intent for {user_type}...")
        
        specific_request = {
            "document_text": parsed_data["document"],
            "user_type": user_type,
            "max_tokens": 150,
            "temperature": 0.7
        }
        
        specific_response = requests.post("http://localhost:8000/api/v1/customer-intent/generate", json=specific_request)
        
        if specific_response.status_code != 200:
            print(f"❌ Failed: {specific_response.status_code}")
            print(specific_response.text)
        else:
            print(f"✅ Successfully generated intent for {user_type}")
            specific_data = specific_response.json()
            
            print(f"\n{user_type.title()} Intent:")
            print(f"{'-'*50}\n{specific_data['intent']}\n{'-'*50}")

if __name__ == "__main__":
    main() 